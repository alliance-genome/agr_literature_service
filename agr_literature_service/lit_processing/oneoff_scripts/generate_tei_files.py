import argparse
import logging
import os
from io import BytesIO
import warnings
import requests
import pdfplumber
from lxml import etree
import xml.etree.ElementTree as ET
from sqlalchemy import text, bindparam
from copy import deepcopy
import tempfile
import camelot
import io
from pypdf import PdfReader
from docx import Document
import pandas as pd

import concurrent.futures
import multiprocessing as mp

from agr_literature_service.lit_processing.utils.sqlalchemy_utils import create_postgres_session
from agr_literature_service.api.crud.referencefile_crud import download_file
from agr_literature_service.api.models import ReferenceModel, ReferencefileModel, CrossReferenceModel
from agr_literature_service.api.routers.okta_utils import OktaAccess

NS = "http://www.tei-c.org/ns/1.0"
ET.register_namespace('', NS)
# Suppress Camelot 'no tables' UserWarnings
warnings.filterwarnings("ignore", category=UserWarning, module="camelot")

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

output_dir = "TEI_FILES/"
MAX_SUPPL_SIZE = 5 * 1024 * 1024  # 5 MB
MAX_SUPPL_PAGE = 20
MAX_MAIN_SIZE = 20 * 1024 * 1024  # 20 MB
MAX_MAIN_PAGE = 50


reference_curies = [
    'AGRKB:101000001159692',
    'AGRKB:101000001065798',
    'AGRKB:101000001065200',
    'AGRKB:101000001063400',
    'AGRKB:101000001058555',
    'AGRKB:101000001052842',
    'AGRKB:101000001052821',
    'AGRKB:101000001047742',
    'AGRKB:101000001045346',
    'AGRKB:101000001044729'
]


def indent(elem, level=0):
    """In-place prettyprint formatter for XML elements."""
    i = "\n" + level * "  "
    if len(elem):
        if not elem.text or not elem.text.strip():
            elem.text = i + "  "
        for child in elem:
            indent(child, level + 1)
        if not child.tail or not child.tail.strip():
            child.tail = i
    if level and (not elem.tail or not elem.tail.strip()):
        elem.tail = i
    return elem


def convert_xlsx_to_tei(file_bytes: bytes, output_path: str):
    """
    Convert an Excel (XLSX) file to a TEI XML document.
    Each sheet becomes a <div type="sheet" xml:id="..."> containing a <table>.
    """
    # Read workbook
    # xl = pd.read_excel(input_path, sheet_name=None)
    xl = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    
    # Create TEI root
    tei = ET.Element("{%(ns)s}TEI" % {'ns': NS})
    text = ET.SubElement(tei, "{%(ns)s}text" % {'ns': NS})
    body = ET.SubElement(text, "{%(ns)s}body" % {'ns': NS})

    base = os.path.splitext(os.path.basename(output_path))[0]

    for idx, (sheet_name, df) in enumerate(xl.items(), start=1):
        div = ET.SubElement(body, f"{{{NS}}}div", {
            'type': 'sheet',
            'xml:id': f"{base}-sheet{idx}"
        })
        head = ET.SubElement(div, f"{{{NS}}}head")
        head.text = sheet_name
        table = ET.SubElement(div, f"{{{NS}}}table")

        # Header row
        if not df.empty:
            hdr = ET.SubElement(table, f"{{{NS}}}row")
            for col in df.columns:
                cell = ET.SubElement(hdr, f"{{{NS}}}cell", {'role': 'header'})
                cell.text = str(col)
            # Data rows
            for _, row in df.iterrows():
                r = ET.SubElement(table, f"{{{NS}}}row")
                for val in row:
                    c = ET.SubElement(r, f"{{{NS}}}cell")
                    c.text = '' if pd.isna(val) else str(val)

    indent(tei)
    tree = ET.ElementTree(tei)
    tree.write(output_path, encoding='utf-8', xml_declaration=True)
    logger.info(f"Converted XLSX bytes → TEI '{output_path}'")


def convert_docx_to_tei(file_bytes: bytes, output_path: str):
    """
    Convert a Word (DOCX) file to a TEI XML document.
    Paragraphs become <p>, headings (Heading 1-6) become <head>.
    """
    # doc = Document(input_path)
    doc = Document(io.BytesIO(file_bytes))
    
    tei = ET.Element("{%(ns)s}TEI" % {'ns': NS})
    text = ET.SubElement(tei, "{%(ns)s}text" % {'ns': NS})
    body = ET.SubElement(text, "{%(ns)s}body" % {'ns': NS})

    for para in doc.paragraphs:
        style = para.style.name.lower()
        if style.startswith('heading'):
            el = ET.SubElement(body, "{%(ns)s}head" % {'ns': NS})
            el.text = para.text
        else:
            el = ET.SubElement(body, "{%(ns)s}p" % {'ns': NS})
            el.text = para.text

    indent(tei)
    tree = ET.ElementTree(tei)
    tree.write(output_path, encoding='utf-8', xml_declaration=True)
    logger.info(f"Converted DOCX bytes → TEI '{output_path}'")


def get_curie_to_reffile_id_mapping(db):

    sql = text("""
        SELECT
          r.curie, rf.referencefile_id, rf.file_class, rf.display_name, rf.file_extension
        FROM reference r
        JOIN referencefile rf
          ON r.reference_id = rf.reference_id
        WHERE r.curie IN :ref_curies
          AND rf.file_extension in ('pdf', 'xlsx', 'docx')
          AND rf.file_publication_status = 'final'
        ORDER BY
          r.curie,
          CASE rf.file_class
            WHEN 'main' THEN 0
            WHEN 'supplement' THEN 1
            ELSE 2
        END
    """).bindparams(bindparam("ref_curies", expanding=True))
    rows = db.execute(sql, {"ref_curies": reference_curies}).all()
    mapping = {}
    for curie, rf_id, file_class, display_name, file_extension in rows:
        mapping.setdefault(curie, []).append((rf_id, file_class, display_name, file_extension))
    return mapping


def convert_pdf_with_grobid(file_content: bytes) -> requests.Response:
    """Send PDF bytes to GROBID and return the HTTP response."""
    url = os.environ.get(
        "PDF2TEI_API_URL",
        "https://grobid.alliancegenome.org/api/processFulltextDocument"
    )
    return requests.post(url, files={"input": ("file", file_content)})


def combine_tei_documents(tei_roots):
    """
    Given a list of parsed TEI roots, produce a single combined TEI tree.
    We merge all <body> children under one master <body>, stripping xml:id and id attributes
    to avoid duplicates and ensure valid TEI.
    Also include document-level <head>, <title>, and <abstract> from the first TEI - from main PDF.
    """
    nsmap = {None: NS}
    master = etree.Element(f"{{{NS}}}TEI", nsmap=nsmap)
    tei_text = etree.SubElement(master, f"{{{NS}}}text")
    # Add front section for title and abstract
    tei_front = etree.SubElement(tei_text, f"{{{NS}}}front")
    # extract title and abstract from the first TEI root
    first_root = tei_roots[0]
    head_el = first_root.find(f".//{{{NS}}}head")
    if head_el is not None:
        tei_front.append(deepcopy(head_el))
    title_el = first_root.find(f".//{{{NS}}}title")
    if title_el is not None:
        tei_front.append(deepcopy(title_el))
    abstract_el = first_root.find(f".//{{{NS}}}abstract")
    if abstract_el is not None:
        tei_front.append(deepcopy(abstract_el))
    # create combined body
    tei_body = etree.SubElement(tei_text, f"{{{NS}}}body")

    for root in tei_roots:
        if root is None:
            continue
        # Remove any xml:id or generic id attributes in this root subtree
        for el in root.iter():
            xml_id_attr = "{http://www.w3.org/XML/1998/namespace}id"
            if xml_id_attr in el.attrib:
                del el.attrib[xml_id_attr]
            if "id" in el.attrib:
                del el.attrib["id"]

        # Append all children of <body> into the combined TEI body
        body = root.find(f".//{{{NS}}}body")
        if body is not None:
            for child in body:
                tei_body.append(child)

    return master


def has_any_table_candidate(pdf_path):
    """Quickly scan PDF for obvious table cues: ruling lines or consistent column-like text spacing."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # 1. If there are ruling lines (used by lattice), that's a strong signal.
                if page.lines:
                    return True
                # 2. Heuristic: detect repeated x-coordinate clusters in characters (simple columnar text)
                chars = page.chars
                if not chars:
                    continue
                # bucket x0 positions into bins to see if there are >=3 repeating columns
                x_buckets = {}
                for ch in chars:
                    x = round(ch.get("x0", 0))
                    x_buckets[x] = x_buckets.get(x, 0) + 1
                # if multiple distinct x positions with decent counts, could be tabular
                large_buckets = [cnt for cnt in x_buckets.values() if cnt >= 10]
                if len(large_buckets) >= 3:
                    return True
    except Exception:
        # If pdfplumber fails for any reason, fall back to running Camelot (safe default)
        return True
    return False


def convert_pdf_to_tei(ref_file_id, reference_curie, file_content, file_class, output_path):
    """Do GROBID + Camelot table extraction for one PDF. Returns (tei_root, table_elements)."""
    tei_root = None
    table_elems = []

    size = len(file_content)
    if (file_class == 'supplement' and size > MAX_SUPPL_SIZE) or (file_class == 'main' and size > MAX_MAIN_SIZE):
        logger.warning(f"{file_class} PDF {ref_file_id} for {reference_curie} too large ({size / (1024*1024):.1f} MB); skipping.")
        return

    reader = PdfReader(io.BytesIO(file_content))
    pages = len(reader.pages)
    if (file_class == 'supplement' and pages > MAX_SUPPL_PAGE) or (file_class == 'main' and pages > MAX_MAIN_PAGE):
        logger.warning(f"{file_class} PDF {ref_file_id} for {reference_curie} too large: total {pages} pages; skipping.")
        return

    # 1) GROBID
    resp = convert_pdf_with_grobid(file_content)
    resp.raise_for_status()
    tei_root = etree.fromstring(resp.content)
    logger.info(f"GROBID converted PDF {ref_file_id} for {reference_curie} file_size={size / (1024*1024):.1f} MB, {pages} pages (class={file_class})")

    # Quick pre-check: if no candidate tables, skip Camelot.
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name

    try:
        if not has_any_table_candidate(tmp_path):
            logger.info(f"No table-like structures detected in {reference_curie}/{ref_file_id}; skipping Camelot.") 
        else:
            # 2) Camelot tables
            try:
                lattice_tables = run_camelot_with_proc_timeout(tmp_path, 'lattice', timeout_sec=60)
            except Exception as e:
                logger.warning(f"Lattice parse failed for {reference_curie}/{ref_file_id}: {e}")
                lattice_tables = []
            try:
                stream_tables  = run_camelot_with_proc_timeout(tmp_path, 'stream', timeout_sec=60)
            except Exception as e:
                logger.warning(f"Stream parse failed for {reference_curie}/{ref_file_id}: {e}")
                stream_tables = []

            # Merge while avoiding duplicates
            tables = list(lattice_tables) + [tbl for tbl in stream_tables if tbl not in lattice_tables]
            for table in tables:
                df = table.df
                tbl_elem = etree.Element(f"{{{NS}}}table", nsmap={None: NS})
                for _, row in df.iterrows():
                    tr = etree.SubElement(tbl_elem, f"{{{NS}}}row")
                    for cell in row:
                        td = etree.SubElement(tr, f"{{{NS}}}cell")
                        td.text = str(cell)
                table_elems.append(tbl_elem)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    body = tei_root.find(f".//{{{NS}}}body")
    for tbl in table_elems:
        body.append(tbl)
    combined_bytes = etree.tostring(tei_root, xml_declaration=True, encoding="utf-8")
    with open(output_path, "wb") as out_f:
        out_f.write(combined_bytes)
    logger.info(f"Converted PDF → TEI '{output_path}'")


def _camelot_worker(path, flavor, conn):
    """Worker run inside its own process; sends back pickled table list or error."""
    try:
        tables = camelot.read_pdf(path, pages="all", flavor=flavor)
        # We don't send the full Camelot Table objects if they aren't picklable; serialize only df
        simplified = [(tbl.df.to_dict(), tbl.shape) for tbl in tables]
        conn.send(("ok", simplified))
    except Exception as e:
        conn.send(("error", str(e)))
    finally:
        conn.close()


def run_camelot_with_proc_timeout(path, flavor, timeout_sec=25):
    parent_conn, child_conn = mp.Pipe()
    proc = mp.Process(target=_camelot_worker, args=(path, flavor, child_conn), daemon=True)
    proc.start()
    proc.join(timeout_sec)
    if proc.is_alive():
        logger.warning(f"Camelot {flavor} timed out on {path}; terminating process.")
        proc.terminate()
        proc.join()
        return []
    if parent_conn.poll():
        status, payload = parent_conn.recv()
        if status == "ok":
            # reconstruct minimal Camelot-like placeholder from simplified data
            tables = []
            for df_dict, shape in payload:
                # rebuild a DataFrame for downstream usage if needed
                import pandas as pd
                df = pd.DataFrame.from_dict(df_dict)
                class DummyTable:
                    def __init__(self, df):
                        self.df = df
                        self.shape = df.shape
                tables.append(DummyTable(df))
            return tables
        else:
            logger.warning(f"Camelot {flavor} failed on {path}: {payload}")
            return []
    else:
        logger.warning(f"No response from Camelot {flavor} worker on {path}.")
        return []


def main():
    parser = argparse.ArgumentParser(description="Combine TEI for references with caching to avoid reprocessing.")
    parser.add_argument("--force", action="store_true", help="Re-generate combined TEI even if it already exists.")
    args = parser.parse_args()

    os.makedirs(output_dir, exist_ok=True)
    db = create_postgres_session(False)
    errors = []

    curie_to_reffile_id_mapping = get_curie_to_reffile_id_mapping(db)

    for reference_curie, ref_ids_with_other_data in curie_to_reffile_id_mapping.items():
        # found_main_pdf = False
        processed_ref_file_ids = set()
        for (ref_file_id, file_class, display_name, file_extension) in ref_ids_with_other_data:
            # only want one main pdf - for this testing purpose 
            if file_class == 'main':
                """
                if found_main_pdf:
                    continue
                found_main_pdf = True
                """
                continue
            if ref_file_id in processed_ref_file_ids:
                continue
            processed_ref_file_ids.add(ref_file_id)

            logger.info(f"{reference_curie}: {file_class} {display_name}.{file_extension}")

            if file_extension not in ['pdf', 'xlsx', 'docx']:
                continue
            
            try:
                # Download file bytes
                file_content = download_file(
                    db=db,
                    referencefile_id=ref_file_id,
                    mod_access=OktaAccess.ALL_ACCESS,
                    use_in_api=False
                )
                os.makedirs(output_dir + f"{reference_curie}", exist_ok=True)
                out_path = os.path.join(output_dir, f"{reference_curie}/{display_name}.tei")
                if file_extension == 'docx':
                    convert_docx_to_tei(file_content, out_path)
                elif file_extension == 'xlsv':
                    convert_xlsx_to_tei(file_content, out_path)
                else:
                    convert_pdf_to_tei(ref_file_id, reference_curie, file_content, file_class, out_path)
            except Exception as e:
                logger.error(f"Error processing PDF {ref_file_id} for {reference_curie}: {e}")
                errors.append((reference_curie, ref_file_id, str(e)))

    db.close()

    if errors:
        logger.warning(f"Completed with {len(errors)} errors; see logs for details")


if __name__ == "__main__":
    main()
